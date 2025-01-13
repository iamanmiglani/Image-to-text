import streamlit as st
import os
import json
from docx import Document
from fpdf import FPDF
import easyocr
import tempfile
import uuid
from PIL import Image
import pyheif
from datetime import datetime, timedelta
import logging

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# File-based lock and queue configuration
LOCK_FILE = "/tmp/app_lock.lock"
QUEUE_FILE = "/tmp/user_queue.json"
LOCK_TIMEOUT = 300  # 5 minutes

def load_queue():
    """Load the queue from a file."""
    if os.path.exists(QUEUE_FILE):
        with open(QUEUE_FILE, "r") as file:
            return json.load(file)
    return []

def save_queue(queue):
    """Save the queue to a file."""
    with open(QUEUE_FILE, "w") as file:
        json.dump(queue, file)

def enqueue_user(user_id):
    """Add a user to the end of the queue."""
    queue = load_queue()
    if user_id not in queue:
        queue.append(user_id)
        save_queue(queue)

def dequeue_user():
    """Remove and return the first user in the queue."""
    queue = load_queue()
    if queue:
        next_user = queue.pop(0)
        save_queue(queue)
        return next_user
    return None

def get_current_user():
    """Get the current user (first in the queue)."""
    queue = load_queue()
    return queue[0] if queue else None

def acquire_lock():
    """Create a lock file if it doesn't exist or belongs to the current session."""
    try:
        if os.path.exists(LOCK_FILE):
            with open(LOCK_FILE, "r") as lock_file:
                lock_time = datetime.fromisoformat(lock_file.read().strip())
                if datetime.now() > lock_time + timedelta(seconds=LOCK_TIMEOUT):
                    # Lock expired, remove it
                    os.remove(LOCK_FILE)
                    logging.debug("Expired lock file removed.")
                else:
                    logging.debug("Lock is active. Another user is using the app.")
                    return False
        with open(LOCK_FILE, "w") as lock_file:
            lock_file.write(datetime.now().isoformat())  # Write the current lock time
            logging.debug("Lock acquired successfully.")
        return True
    except Exception as e:
        logging.error(f"Error in acquire_lock: {e}")
        raise

def release_lock():
    """Remove the lock file."""
    try:
        if os.path.exists(LOCK_FILE):
            os.remove(LOCK_FILE)
    except Exception as e:
        logging.error(f"Error in release_lock: {e}")
        raise

# Preload EasyOCR reader
@st.cache_resource
def load_easyocr_reader():
    return easyocr.Reader(["en"], gpu=False)

def convert_heic_to_png(image_file):
    heif_file = pyheif.read(image_file.read())
    image = Image.frombytes(
        heif_file.mode, heif_file.size, heif_file.data, "raw", heif_file.mode, heif_file.stride
    )
    return image

def extract_text_from_images(images, reader):
    extracted_text = {}
    for image_file in images:
        if image_file.type == "image/heic":
            image = convert_heic_to_png(image_file)
            image_bytes = image.tobytes()
        else:
            image_bytes = image_file.read()

        results = reader.readtext(image_bytes, detail=0)
        extracted_text[image_file.name] = results
    return extracted_text

def generate_word_document(extracted_text):
    doc = Document()
    doc.add_heading("Extracted Text from Images", level=1)
    for image_name, text in extracted_text.items():
        doc.add_heading(f"Image: {image_name}", level=2)
        for line in text:
            doc.add_paragraph(line)
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_extracted_text.docx")
    doc.save(output_path)
    return output_path

def generate_pdf_document(extracted_text):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Extracted Text from Images", ln=True, align='C')
    pdf.ln(10)
    for image_name, text in extracted_text.items():
        pdf.set_font("Arial", style='B', size=12)
        pdf.cell(200, 10, txt=f"Image: {image_name}", ln=True)
        pdf.set_font("Arial", size=12)
        for line in text:
            pdf.multi_cell(0, 10, txt=line)
        pdf.ln(5)
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_extracted_text.pdf")
    pdf.output(output_path)
    return output_path

def reset_session():
    """Clear all session variables and reload the app."""
    release_lock()
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.experimental_rerun()

def main():
    if "user_id" not in st.session_state:
        # Unique session identifier (e.g., user IP or session ID)
        st.session_state.user_id = str(uuid.uuid4())  # Generate unique user ID

    # Add the user to the queue if not already present
    enqueue_user(st.session_state.user_id)

    # Check if the current user is at the front of the queue
    if st.session_state.user_id != get_current_user():
        st.warning("The app is currently in use by another user. Please wait for your turn.")
        st.stop()

    # User is at the front of the queue
    st.title("Image Text Extraction App")

    # Load EasyOCR reader (cached for performance)
    reader = load_easyocr_reader()

    if "extracted_text" not in st.session_state:
        st.session_state.extracted_text = None
    if "file_path" not in st.session_state:
        st.session_state.file_path = None
    if "download_complete" not in st.session_state:
        st.session_state.download_complete = False

    if not st.session_state.download_complete:
        uploaded_files = st.file_uploader(
            "Upload Images (Max 10)", type=["jpg", "jpeg", "png", "heic"], accept_multiple_files=True
        )

        if uploaded_files:
            if st.session_state.extracted_text is None:
                with st.spinner("Extracting text..."):
                    if len(uploaded_files) > 10:
                        st.error("You can upload a maximum of 10 images.")
                    else:
                        st.session_state.extracted_text = extract_text_from_images(uploaded_files, reader)
                        st.success("Text extraction complete!")

        if st.session_state.extracted_text:
            output_format = st.selectbox("Select Output Format", ["Word", "PDF"])
            if output_format:
                st.session_state.output_format = output_format

            if st.button("Generate Document"):
                with st.spinner("Preparing your document..."):
                    if st.session_state.output_format == "Word":
                        file_path = generate_word_document(st.session_state.extracted_text)
                    else:
                        file_path = generate_pdf_document(st.session_state.extracted_text)

                    st.session_state.file_path = file_path
                    st.success(f"{st.session_state.output_format} document ready!")

            if st.session_state.file_path:
                with open(st.session_state.file_path, "rb") as file:
                    download_button_clicked = st.download_button(
                        label=f"Download {st.session_state.output_format} Document",
                        data=file,
                        file_name=os.path.basename(st.session_state.file_path),
                        mime="application/octet-stream",
                    )

                    if download_button_clicked:
                        st.session_state.download_complete = True
                        st.experimental_rerun()
    else:
        st.info("Do you want to use the app again?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes, Reset"):
                reset_session()  # Clear session and restart
        with col2:
            if st.button("No, Exit"):
                st.success("Thanks for using the app!")
                release_lock()  # Release the lock for others
                enqueue_user(st.session_state.user_id)  # Add user to the end of the queue
                dequeue_user()  # Allow the next user in the queue
                st.stop()  # Stop the app completely

if __name__ == "__main__":
    main()
